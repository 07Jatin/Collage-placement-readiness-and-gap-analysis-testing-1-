from pathlib import Path
from math import ceil

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "generated"
IMG_DIR = OUT_DIR / "diagram_assets"
OUT_DOCX = OUT_DIR / "College_Placement_Readiness_and_Gap_Analysis_Documentation.docx"

BLUE = "1F4E79"
NAVY = "0F2437"
TEAL = "0F766E"
GREEN = "2E7D32"
GOLD = "B7791F"
RED = "B91C1C"
GRAY = "4B5563"
LIGHT_BLUE = "EAF4FB"
LIGHT_TEAL = "E8F7F5"
LIGHT_GREEN = "ECF7EC"
LIGHT_GOLD = "FFF7E6"
LIGHT_RED = "FDECEC"
LIGHT_GRAY = "F3F4F6"
WHITE = "FFFFFF"


def font(size=26, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
        "arial.ttf",
    ]
    for item in candidates:
        try:
            return ImageFont.truetype(item, size)
        except Exception:
            pass
    return ImageFont.load_default()


def wrap_text(draw, text, fnt, max_width):
    words = str(text).split()
    lines = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [""]


def draw_box(draw, xy, text, fill, outline, text_color=NAVY, radius=18, font_size=24, bold=True):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=f"#{fill}", outline=f"#{outline}", width=3)
    fnt = font(font_size, bold)
    lines = wrap_text(draw, text, fnt, x2 - x1 - 28)
    line_h = fnt.size + 7
    total_h = line_h * len(lines)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, font=fnt, fill=f"#{text_color}")
        y += line_h


def arrow(draw, p1, p2, color=GRAY, width=4):
    draw.line([p1, p2], fill=f"#{color}", width=width)
    x1, y1 = p1
    x2, y2 = p2
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 >= x1 else -1
        pts = [(x2, y2), (x2 - 16 * direction, y2 - 10), (x2 - 16 * direction, y2 + 10)]
    else:
        direction = 1 if y2 >= y1 else -1
        pts = [(x2, y2), (x2 - 10, y2 - 16 * direction), (x2 + 10, y2 - 16 * direction)]
    draw.polygon(pts, fill=f"#{color}")


def make_linear_diagram(path, title, nodes, direction="vertical", accent=BLUE):
    w, h = (1600, max(650, len(nodes) * 135 + 170)) if direction == "vertical" else (1900, 650)
    img = Image.new("RGB", (w, h), f"#{WHITE}")
    d = ImageDraw.Draw(img)
    d.text((60, 40), title, font=font(42, True), fill=f"#{NAVY}")
    if direction == "vertical":
        box_w, box_h = 760, 78
        x = (w - box_w) // 2
        y = 135
        palette = [LIGHT_BLUE, LIGHT_TEAL, LIGHT_GREEN, LIGHT_GOLD]
        for i, node in enumerate(nodes):
            draw_box(d, (x, y, x + box_w, y + box_h), node, palette[i % len(palette)], accent, font_size=23)
            if i < len(nodes) - 1:
                arrow(d, (x + box_w // 2, y + box_h + 4), (x + box_w // 2, y + box_h + 45), accent)
            y += 125
    else:
        cols = min(5, len(nodes))
        rows = ceil(len(nodes) / cols)
        box_w, box_h = 300, 88
        gap_x, gap_y = 56, 90
        start_x = (w - (cols * box_w + (cols - 1) * gap_x)) // 2
        start_y = 150
        for idx, node in enumerate(nodes):
            row, col = divmod(idx, cols)
            if row % 2 == 1:
                col = cols - 1 - col
            x = start_x + col * (box_w + gap_x)
            y = start_y + row * (box_h + gap_y)
            draw_box(d, (x, y, x + box_w, y + box_h), node, LIGHT_BLUE if idx % 2 == 0 else LIGHT_TEAL, accent, font_size=21)
            if idx < len(nodes) - 1:
                nrow, ncol = divmod(idx + 1, cols)
                if nrow % 2 == 1:
                    ncol = cols - 1 - ncol
                nx = start_x + ncol * (box_w + gap_x)
                ny = start_y + nrow * (box_h + gap_y)
                if row == nrow:
                    start = (x + box_w + 4, y + box_h // 2) if nx > x else (x - 4, y + box_h // 2)
                    end = (nx - 4, ny + box_h // 2) if nx > x else (nx + box_w + 4, ny + box_h // 2)
                    arrow(d, start, end, accent)
                else:
                    arrow(d, (x + box_w // 2, y + box_h + 4), (x + box_w // 2, ny - 18), accent)
                    arrow(d, (x + box_w // 2, ny - 18), (nx + box_w // 2, ny - 18), accent)
                    arrow(d, (nx + box_w // 2, ny - 18), (nx + box_w // 2, ny - 4), accent)
    img.save(path)


def make_architecture(path):
    img = Image.new("RGB", (1800, 1050), f"#{WHITE}")
    d = ImageDraw.Draw(img)
    d.text((60, 42), "Application Architecture", font=font(46, True), fill=f"#{NAVY}")
    boxes = {
        "User Browser / Electron Shell": (110, 170, 480, 260, LIGHT_BLUE, BLUE),
        "React Frontend": (690, 170, 1080, 260, LIGHT_TEAL, TEAL),
        "FastAPI Backend": (690, 430, 1080, 520, LIGHT_BLUE, BLUE),
        "Business Logic": (360, 680, 690, 770, LIGHT_GOLD, GOLD),
        "Readiness Rules": (750, 680, 1080, 770, LIGHT_GREEN, GREEN),
        "Gap Analysis": (1140, 680, 1470, 770, LIGHT_RED, RED),
        "Local JSON / CSV Data": (1230, 430, 1620, 520, LIGHT_GRAY, GRAY),
    }
    for text, (x1, y1, x2, y2, fill, outline) in boxes.items():
        draw_box(d, (x1, y1, x2, y2), text, fill, outline, font_size=24)
    arrow(d, (480, 215), (690, 215), BLUE)
    arrow(d, (885, 260), (885, 430), TEAL)
    arrow(d, (1080, 475), (1230, 475), BLUE)
    arrow(d, (885, 520), (525, 680), BLUE)
    arrow(d, (885, 520), (915, 680), GREEN)
    arrow(d, (885, 520), (1305, 680), RED)
    arrow(d, (1230, 475), (1080, 475), GRAY)
    arrow(d, (690, 475), (480, 215), BLUE)
    img.save(path)


def make_sequence(path):
    img = Image.new("RGB", (1800, 950), f"#{WHITE}")
    d = ImageDraw.Draw(img)
    d.text((60, 42), "Authentication Flow", font=font(46, True), fill=f"#{NAVY}")
    actors = ["User", "React Frontend", "FastAPI Backend", "Session Store"]
    xs = [230, 650, 1080, 1500]
    for x, actor in zip(xs, actors):
        draw_box(d, (x - 145, 150, x + 145, 220), actor, LIGHT_BLUE, BLUE, font_size=23)
        d.line([(x, 230), (x, 840)], fill="#9CA3AF", width=3)
    steps = [
        (0, 1, "Enter username / roll no and password"),
        (1, 2, "POST /api/login or /api/students/login"),
        (2, 3, "Validate hash and create token"),
        (3, 2, "Return session token"),
        (2, 1, "Send role and token"),
        (1, 0, "Redirect to dashboard"),
    ]
    y = 285
    for src, dst, label in steps:
        x1, x2 = xs[src], xs[dst]
        arrow(d, (x1, y), (x2, y), BLUE if src < dst else TEAL)
        d.text((min(x1, x2) + 25, y - 36), label, font=font(20), fill=f"#{GRAY}")
        y += 86
    img.save(path)


def make_erd(path):
    img = Image.new("RGB", (1800, 1050), f"#{WHITE}")
    d = ImageDraw.Draw(img)
    d.text((60, 42), "Data Model - Conceptual ERD", font=font(46, True), fill=f"#{NAVY}")
    entities = {
        "USER": (130, 190, ["id", "name", "email", "password_hash", "role"]),
        "STUDENT_PROFILE": (585, 190, ["id", "user_id", "department", "year", "cgpa", "target_role"]),
        "ASSESSMENT": (1040, 190, ["id", "student_id", "skill_scores", "test_history", "created_at"]),
        "REPORT": (585, 610, ["id", "student_id", "readiness_score", "gaps", "recommendations"]),
        "MARKET_DATA": (1040, 610, ["role", "required_skills", "skill_weight", "job_trends"]),
    }
    for name, (x, y, fields) in entities.items():
        draw_box(d, (x, y, x + 330, y + 70), name, LIGHT_BLUE, BLUE, font_size=24)
        d.rounded_rectangle((x, y + 70, x + 330, y + 260), radius=10, fill=f"#{WHITE}", outline=f"#{BLUE}", width=3)
        yy = y + 92
        for field in fields:
            d.text((x + 24, yy), field, font=font(21), fill=f"#{NAVY}")
            yy += 32
    arrow(d, (460, 320), (585, 320), TEAL)
    d.text((478, 285), "has", font=font(20, True), fill=f"#{TEAL}")
    arrow(d, (915, 320), (1040, 320), TEAL)
    d.text((933, 285), "submits", font=font(20, True), fill=f"#{TEAL}")
    arrow(d, (1205, 450), (915, 610), TEAL)
    d.text((1020, 518), "generates", font=font(20, True), fill=f"#{TEAL}")
    arrow(d, (1205, 610), (1205, 450), GOLD)
    d.text((1230, 530), "references", font=font(20, True), fill=f"#{GOLD}")
    img.save(path)


def make_score(path):
    img = Image.new("RGB", (1600, 850), f"#{WHITE}")
    d = ImageDraw.Draw(img)
    d.text((60, 42), "Readiness Score Engine", font=font(46, True), fill=f"#{NAVY}")
    center = (800, 440)
    factors = [
        ("Technical\n30%", 800, 155, LIGHT_BLUE, BLUE),
        ("Aptitude\n20%", 1120, 280, LIGHT_TEAL, TEAL),
        ("Communication\n15%", 1060, 610, LIGHT_GREEN, GREEN),
        ("Projects\n15%", 540, 610, LIGHT_GOLD, GOLD),
        ("Resume\n10%", 480, 280, LIGHT_RED, RED),
        ("Interview\n10%", 800, 715, LIGHT_GRAY, GRAY),
    ]
    draw_box(d, (610, 350, 990, 520), "Overall Placement\nReadiness Score", LIGHT_BLUE, BLUE, font_size=30)
    for text, x, y, fill, outline in factors:
        draw_box(d, (x - 145, y - 55, x + 145, y + 55), text.replace("\n", " "), fill, outline, font_size=22)
        arrow(d, (x, y + (55 if y < center[1] else -55)), center, outline)
    img.save(path)


def make_diagrams():
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = {}
    diagrams["flow"] = IMG_DIR / "01_application_flow.png"
    make_linear_diagram(diagrams["flow"], "Basic Web Application Flow", [
        "User opens website",
        "Login or register",
        "Student dashboard",
        "Fill profile details",
        "Enter skills and academic information",
        "Take assessment",
        "Calculate readiness score",
        "Gap analysis report",
        "Personalized recommendations",
        "Improve skills and reassess",
    ], "vertical", BLUE)
    diagrams["roadmap"] = IMG_DIR / "02_project_roadmap.png"
    make_linear_diagram(diagrams["roadmap"], "High-Level Project Roadmap", [
        "Project idea", "Requirement analysis", "UI/UX design", "Frontend development",
        "Backend development", "Data integration", "Readiness logic", "Gap module",
        "Testing", "Deployment", "Future enhancements"
    ], "horizontal", TEAL)
    diagrams["architecture"] = IMG_DIR / "03_architecture.png"
    make_architecture(diagrams["architecture"])
    diagrams["auth"] = IMG_DIR / "04_authentication_flow.png"
    make_sequence(diagrams["auth"])
    diagrams["erd"] = IMG_DIR / "05_database_erd.png"
    make_erd(diagrams["erd"])
    diagrams["score"] = IMG_DIR / "06_score_engine.png"
    make_score(diagrams["score"])
    diagrams["beginner"] = IMG_DIR / "07_beginner_roadmap.png"
    make_linear_diagram(diagrams["beginner"], "Beginner Roadmap", [
        "Learn HTML, CSS, JavaScript",
        "Understand React components",
        "Forms and state management",
        "API calls with fetch or Axios",
        "Backend basics",
        "FastAPI routes and services",
        "Data files and models",
        "Authentication",
        "Readiness score logic",
        "Gap analysis",
        "Test and deploy",
    ], "vertical", GREEN)
    diagrams["developer"] = IMG_DIR / "08_developer_roadmap.png"
    make_linear_diagram(diagrams["developer"], "Developer Roadmap", [
        "Create project setup", "Build UI views", "Create backend server", "Connect local data",
        "Create auth APIs", "Create student APIs", "Assessment logic", "Readiness score",
        "Gap report", "Charts and dashboard", "Test", "Package or deploy"
    ], "horizontal", GOLD)
    return diagrams


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=NAVY, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    r.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    r.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header_fill=BLUE):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if i == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True
            else:
                set_cell_shading(cell, "FFFFFF" if i % 2 else "F8FAFC")
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.space_before = Pt(3)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(BLUE if level == 1 else TEAL)
        run.font.name = "Aptos Display"
    return p


def add_para(doc, text, bold_start=None):
    p = doc.add_paragraph()
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(NAVY)
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.08
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(3)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(3)


def add_callout(doc, title, body, fill="EAF4FB"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run("\n" + body)
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(4)
    doc.add_paragraph()


def add_figure(doc, path, caption, width=6.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.color.rgb = RGBColor.from_string(GRAY)
        run.font.size = Pt(9)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=WHITE, size=9)
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            set_cell_text(cells[i], val, size=8.5)
    style_table(table)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def code_block(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(8)
    return p


def build_doc():
    diagrams = make_diagrams()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].font.color.rgb = RGBColor.from_string(NAVY)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Aptos Display"
        styles[style_name].font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Placify AI")
    r.bold = True
    r.font.size = Pt(31)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("College Placement Readiness and Gap Analysis")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Web Application Project Documentation")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    add_callout(
        doc,
        "Project Snapshot",
        "Repository: https://github.com/07Jatin/Collage-placement-readiness-and-gap-analysis-testing-1-\n"
        "Project type: React + FastAPI web application packaged with Electron\n"
        "Purpose: Analyze student placement readiness, identify skill gaps, and recommend improvement areas.",
        "EAF4FB",
    )
    add_figure(doc, diagrams["architecture"], "Figure 1. Real project architecture overview", width=6.4)
    doc.add_page_break()

    add_heading(doc, "1. Short Introduction")
    add_para(doc, "Placify AI is a placement readiness and gap analysis platform for students, faculty, and placement teams. It helps a student understand how prepared they are for campus placements by combining profile data, test performance, resume analysis, skill-gap checks, and market-aligned role expectations.")
    add_bullets(doc, [
        "Technical skills such as programming, data structures, DBMS, web development, and core computer science.",
        "Aptitude, reasoning, English, and communication readiness.",
        "Resume readiness, project quality, internship exposure, and interview preparation.",
        "Role-specific missing skills and personalized learning paths.",
    ])
    code_block(doc, "Overall Readiness: 68%\nStrengths: Programming, Projects\nWeak Areas: Aptitude, Communication\nSuggested Actions: Practice aptitude daily, attend mock interviews")

    add_heading(doc, "2. Problem Statement")
    add_para(doc, "Many students prepare for placements without a clear view of their current readiness. They may solve questions, build projects, or update resumes, but still not know which areas are blocking their placement success.")
    add_bullets(doc, [
        "Students do not always know which skills companies expect.",
        "Weak areas are often discovered late, usually during interviews or mock tests.",
        "Preparation is not always structured around target roles.",
        "Colleges need a batch-level view of readiness and common gaps.",
    ])

    add_heading(doc, "3. Main Objective")
    add_para(doc, "The objective is to provide a web-based system that collects student data, analyzes current preparation, calculates a readiness score, identifies gaps, and recommends a practical learning path.")
    add_numbered(doc, [
        "Collect student profile, skill, resume, and assessment information.",
        "Analyze preparation against placement and role requirements.",
        "Calculate a placement readiness score.",
        "Identify weak or missing skills.",
        "Suggest resources, quizzes, coding problems, and learning paths.",
        "Support admin views for batch-level analysis.",
    ])

    add_heading(doc, "4. Users of the Application")
    add_table(doc, ["User", "What They Can Do"], [
        ("Students", "Login, maintain profile, upload resume text, take tests, view readiness score, inspect gaps, and follow learning recommendations."),
        ("Admin / Faculty", "Login securely, view students, identify common skill gaps, track at-risk students, and review readiness trends."),
        ("Training and Placement Cell", "Plan workshops, prioritize training areas, and use reports for placement-preparation decisions."),
    ], [1.6, 4.8])

    add_heading(doc, "5. Basic Web Application Flow")
    add_figure(doc, diagrams["flow"], "Figure 2. Student journey from login to improvement cycle", width=5.8)

    add_heading(doc, "6. High-Level Roadmap")
    add_figure(doc, diagrams["roadmap"], "Figure 3. Roadmap from idea to deployment", width=6.6)

    add_heading(doc, "7. Beginner-Friendly Explanation")
    add_para(doc, "Think of the project as three cooperating parts: frontend, backend, and data storage. The frontend displays the screens, the backend performs logic, and the local JSON/CSV data files store project data for this implementation.")
    add_table(doc, ["Part", "Actual Project Technology", "Responsibility"], [
        ("Frontend", "React, TailwindCSS, Recharts, lucide-react", "Pages, dashboards, charts, forms, login screens, report views, and resume upload UI."),
        ("Backend", "Python FastAPI, Uvicorn", "API routes, authentication, readiness rules, skill analysis, resume parsing, quiz validation, and coding lab services."),
        ("Data", "JSON and CSV files", "Student data, market data, test history, skill quizzes, and historical placement data."),
        ("Desktop Packaging", "Electron", "Runs the React application as a desktop-style app."),
    ], [1.25, 2.0, 3.15])

    add_heading(doc, "8. Expected Features")
    add_table(doc, ["Feature", "Description"], [
        ("Student login", "Students can log in using roll number and password. Admin users use a separate username/password flow."),
        ("Student dashboard", "Shows profile, readiness summary, gaps, recommendations, and progress signals."),
        ("Profile management", "Stores name, email, mobile, target role, academic information, skills, and related details."),
        ("Resume analyzer", "Extracts skills from resume text using regex and optional LLM-based semantic parsing."),
        ("Gap analysis", "Compares student skills against market or role requirements and returns missing skills."),
        ("Readiness prediction", "Builds a readiness prediction using student profile, test history, and rule-based scoring."),
        ("Learning path", "Recommends resources, videos, quizzes, and improvement actions."),
        ("Mock tests", "Tracks category scores and assessment history."),
        ("DSA coding lab", "Provides static and dynamic coding problems, syntax checks, and compiler execution through external APIs."),
        ("Admin analytics", "Lists common gaps and at-risk students for faculty or placement-cell action."),
    ], [1.65, 4.75])

    add_heading(doc, "9. Admin Features")
    add_bullets(doc, [
        "Secure admin login through /api/login.",
        "View all students through /api/students.",
        "Analyze common skill gaps through /admin/skill_gaps.",
        "Find at-risk students through /admin/at_risk_students.",
        "Use aggregate information to plan training interventions.",
    ])

    add_heading(doc, "10. Actual Technology Stack")
    add_table(doc, ["Layer", "Technology"], [
        ("Frontend", "React 18, TailwindCSS, Recharts, lucide-react, mammoth, pdfjs-dist"),
        ("Backend", "Python, FastAPI, Uvicorn, httpx"),
        ("AI / NLP", "Regex skill extractor, optional Hugging Face / Mistral-style LLM integration"),
        ("External APIs", "YouTube Search API fallback, Adzuna jobs API, LeetCode GraphQL, Piston compiler API"),
        ("Desktop", "Electron, electron-builder"),
        ("Data", "Local JSON and CSV files"),
        ("Tools", "npm, pip, VS Code, Git, GitHub"),
    ], [1.65, 4.75])

    add_heading(doc, "11. Application Architecture")
    add_para(doc, "The React frontend communicates with a FastAPI backend. The backend delegates analysis to service modules such as skill_analyzer, resume_parser, readiness_rules, llm_resume_analyzer, and adzuna_client. Data is loaded from local JSON and CSV files.")
    add_figure(doc, diagrams["architecture"], "Figure 4. Frontend, backend, services, and data relationship", width=6.6)

    add_heading(doc, "12. Folder Structure")
    code_block(doc, """final project/
|-- backend/
|   |-- data/
|   |-- services/
|   |-- tests/
|   `-- main.py
|-- frontend/
|   |-- public/
|   |-- src/
|   |   |-- assets/
|   |   |-- components/
|   |   |-- data/
|   |   |-- utils/
|   |   |-- App.js
|   |   |-- index.js
|   |   `-- index.css
|   |-- package.json
|   |-- tailwind.config.js
|   `-- postcss.config.js
|-- docs/
|-- main.js
|-- preload.js
|-- package.json
`-- requirements.txt""")

    add_heading(doc, "13. Important Folders")
    add_table(doc, ["Folder / File", "Purpose"], [
        ("frontend/src/components", "Contains major React views such as DashboardView, GapAnalysisView, MockTestView, ProfileView, ResumeUploadView, and AdminDashboardView."),
        ("frontend/src/data", "Stores frontend skill configuration such as skillData.js."),
        ("backend/services", "Contains core service logic for skill analysis, readiness prediction, resume parsing, LLM analysis, and external jobs data."),
        ("backend/data", "Stores student data, market data, test history, placement history, and quiz data."),
        ("backend/main.py", "FastAPI application entry point and API route definitions."),
        ("main.js / preload.js", "Electron shell integration for desktop app packaging."),
        ("docs", "Project documentation, diagrams, notebooks, reports, and presentations."),
    ], [2.1, 4.3])

    add_heading(doc, "14. Core Modules")
    add_heading(doc, "14.1 Authentication Module", 2)
    add_para(doc, "The backend includes separate flows for admin login and student login. Admin users are checked against hashed passwords, while student login upserts or validates student access through readiness_rules.")
    add_figure(doc, diagrams["auth"], "Figure 5. Login and token-based session flow", width=6.6)
    add_heading(doc, "14.2 Student Profile Module", 2)
    add_bullets(doc, ["Stores roll number, name, email, mobile, and target role.", "Feeds readiness prediction and role-based recommendations.", "Can be updated through /api/students/update."])
    add_heading(doc, "14.3 Assessment and Mock Test Module", 2)
    add_bullets(doc, ["Skill quizzes are loaded from backend data files.", "Test result submissions are saved with overall and category scores.", "History can be retrieved using /api/test_history/{student_id}."])
    add_heading(doc, "14.4 Readiness Score Module", 2)
    add_para(doc, "The readiness prediction module combines profile information, verified skills, skill gaps, and test history. A simple explainable formula can be used as a beginner-friendly model:")
    code_block(doc, "Readiness Score = Technical * 0.30 + Aptitude * 0.20 + Communication * 0.15 + Projects * 0.15 + Resume * 0.10 + Interview * 0.10")
    add_figure(doc, diagrams["score"], "Figure 6. Weighted readiness scoring concept", width=6.2)
    add_heading(doc, "14.5 Gap Analysis Module", 2)
    add_para(doc, "Gap analysis compares a student's current skills with role or market requirements. Missing skills are returned as a report and also aggregated for admin analytics.")
    add_table(doc, ["Skill Area", "Required", "Student", "Gap"], [
        ("Technical", "75", "80", "No gap"),
        ("Aptitude", "70", "60", "Gap found"),
        ("Communication", "70", "65", "Gap found"),
        ("Resume", "80", "90", "No gap"),
        ("Projects", "70", "55", "Gap found"),
    ], [2.1, 1.1, 1.1, 2.1])
    add_heading(doc, "14.6 Recommendation Module", 2)
    add_table(doc, ["Weak Area", "Recommendation"], [
        ("DSA", "Solve coding problems daily and practice topic-wise LeetCode-style questions."),
        ("Aptitude", "Practice number systems, percentages, profit/loss, reasoning, and timed quizzes."),
        ("Communication", "Join group discussions, record answers, and attend mock interviews."),
        ("Resume", "Add measurable project outcomes, skills, links, and certification details."),
        ("Projects", "Build full-stack or role-specific projects with clear documentation."),
        ("Interview", "Revise HR questions, technical fundamentals, and previous test mistakes."),
    ], [1.55, 4.85])

    add_heading(doc, "15. Database Design")
    add_para(doc, "This project currently uses local JSON and CSV files, but the same design can be moved into MongoDB, PostgreSQL, MySQL, or Supabase later.")
    add_figure(doc, diagrams["erd"], "Figure 7. Conceptual data model for future database migration", width=6.6)

    add_heading(doc, "16. API Documentation")
    add_table(doc, ["Method", "Endpoint", "Purpose"], [
        ("POST", "/api/login", "Admin login, returns token and role."),
        ("POST", "/api/students/login", "Student login or account setup by roll number."),
        ("GET", "/api/students", "Admin-protected student listing."),
        ("POST", "/api/students/update", "Update student profile and target role."),
        ("GET", "/gap_report/{student_id}", "Return skill gap analysis for a student."),
        ("GET", "/predict_readiness/{student_id}", "Return readiness prediction."),
        ("POST", "/parse_resume", "Extract skills from resume text."),
        ("GET", "/skill_quiz/{skill_name}", "Return quiz questions for a skill."),
        ("POST", "/validate_skills", "Validate quiz results and derive verified skills."),
        ("POST", "/api/submit_test_result", "Save mock test result and category scores."),
        ("GET", "/api/test_history/{student_id}", "Return previous test results."),
        ("GET", "/admin/skill_gaps", "Admin aggregate missing skill frequencies."),
        ("GET", "/admin/at_risk_students", "Admin list of students who need support."),
        ("GET", "/api/live_jobs/{role}", "Fetch live job data through Adzuna."),
        ("GET", "/api/youtube/search", "Fetch or mock YouTube learning resources."),
        ("GET", "/api/dsa/dynamic/{skill}", "Fetch a dynamic LeetCode-style problem."),
        ("POST", "/compile", "Compile or check submitted code."),
    ], [0.7, 2.25, 3.45])

    add_heading(doc, "17. Frontend Pages and Views")
    add_table(doc, ["View Component", "Role in Application"], [
        ("LoginView.js", "Handles user login experience."),
        ("DashboardView.js", "Student dashboard summary."),
        ("ProfileView.js", "Student profile and target role information."),
        ("GapAnalysisView.js", "Skill gap report and recommended actions."),
        ("LearningPathView.js", "Learning recommendations and resources."),
        ("MockTestView.js", "Assessment and mock-test experience."),
        ("ResumeUploadView.js", "Resume parsing and skill extraction workflow."),
        ("CodeEditorView.js", "DSA coding lab interface."),
        ("AdminDashboardView.js", "Admin summary and analytics."),
        ("ManageStudentsView.js", "Student management for admin/faculty."),
    ], [2.1, 4.3])

    add_heading(doc, "18. Readiness Status Levels")
    add_table(doc, ["Score Range", "Status", "Meaning"], [
        ("0-40", "Not Ready", "Needs major improvement."),
        ("41-60", "Basic Readiness", "Some preparation done, but important gaps remain."),
        ("61-75", "Moderately Ready", "Good base, still needs targeted practice."),
        ("76-90", "Placement Ready", "Strong preparation for many placement processes."),
        ("91-100", "Highly Ready", "Excellent placement readiness."),
    ], [1.2, 1.7, 3.5])

    add_heading(doc, "19. Example Gap Analysis Output")
    code_block(doc, """Overall Score: 67%
Strengths:
- Technical Skills
- Projects
Weak Areas:
- Aptitude
- Communication
- Interview Preparation
Recommendations:
- Practice aptitude questions for 30 minutes daily.
- Participate in group discussions.
- Attend mock technical and HR interviews.""")

    add_heading(doc, "20. How the Project Is Built")
    add_numbered(doc, [
        "Requirement analysis: decide student, admin, assessment, report, and recommendation features.",
        "UI design: create login, dashboard, profile, assessment, report, resume, and admin views.",
        "Frontend development: build React components and connect them to APIs.",
        "Backend development: define FastAPI routes and service modules.",
        "Data integration: load and update JSON/CSV files.",
        "Readiness logic: calculate readiness using rules, tests, and skill evidence.",
        "Gap analysis: compare current skills with target-role requirements.",
        "Testing: run backend tests, UI checks, and manual API checks.",
        "Packaging/deployment: run as a web app or package with Electron.",
    ])

    add_heading(doc, "21. Testing Strategy")
    add_table(doc, ["Testing Type", "What to Check"], [
        ("Manual testing", "Login, dashboard, profile update, resume parsing, quiz submission, score display, and report pages."),
        ("Backend tests", "FastAPI endpoints, skill analysis, readiness prediction, submit-test flow, and external API fallbacks."),
        ("Frontend testing", "Forms, chart rendering, navigation, validation, loading states, and error messages."),
        ("API testing", "Use Postman or curl for /api/login, /gap_report, /predict_readiness, /parse_resume, and /compile."),
        ("Scoring tests", "Confirm known sample inputs produce expected readiness outputs."),
    ], [1.7, 4.7])

    add_heading(doc, "22. Security Considerations")
    add_bullets(doc, [
        "Avoid storing real passwords as plain text. Use strong password hashing in production.",
        "Move hard-coded admin users to a secure database or environment-managed identity store.",
        "Use persistent session storage such as Redis or a database in production.",
        "Protect admin routes with bearer tokens and role checks.",
        "Keep API keys such as YouTube, Adzuna, and Hugging Face tokens in environment variables.",
        "Validate inputs for resume text, code submissions, profile updates, and generated question requests.",
    ])

    add_heading(doc, "23. Deployment Process")
    add_table(doc, ["Part", "Command / Platform"], [
        ("Backend local", "cd backend; uvicorn main:app --reload --port 8000"),
        ("Frontend local", "cd frontend; npm install; npm start"),
        ("Electron local", "npm install; npm run dev"),
        ("Frontend deployment", "Build with npm run build --prefix frontend, then deploy to Vercel or Netlify."),
        ("Backend deployment", "Deploy FastAPI to Render, Railway, Azure App Service, or similar."),
        ("Future database", "Move JSON/CSV storage to MongoDB Atlas, PostgreSQL, Supabase, or MySQL."),
    ], [1.8, 4.6])

    add_heading(doc, "24. Beginner Roadmap")
    add_figure(doc, diagrams["beginner"], "Figure 8. Recommended learning order for beginners", width=5.8)

    add_heading(doc, "25. Developer Roadmap")
    add_figure(doc, diagrams["developer"], "Figure 9. Practical build roadmap for developers", width=6.6)

    add_heading(doc, "26. Advanced Improvements")
    add_table(doc, ["Improvement", "Benefit"], [
        ("AI-based recommendations", "Generate personalized 30-day plans based on gaps and target roles."),
        ("Resume analyzer upload", "Support direct PDF/DOCX upload with deeper ATS feedback."),
        ("Company-specific readiness", "Compare students against TCS, Infosys, Wipro, Amazon, or other company patterns."),
        ("Learning resources", "Attach curated videos, notes, quizzes, and coding tasks to every weak area."),
        ("Progress tracking", "Show trend charts across repeated assessments."),
        ("Database migration", "Replace local files with a production database and audit trail."),
        ("Role-based admin dashboard", "Separate faculty, placement-cell, and super-admin permissions."),
    ], [2.1, 4.3])

    add_heading(doc, "27. Example Final Report Format")
    code_block(doc, """Student Name: Rahul Sharma
Department: Computer Science
Year: 4th Year
Overall Readiness Score: 72%
Status: Moderately Ready
Strengths: Technical Skills, Resume, Projects
Weak Areas: Aptitude, Interview Preparation, Communication
Recommendations:
1. Practice aptitude questions daily.
2. Attend weekly mock interviews.
3. Improve communication through group discussions.
4. Revise DBMS, OOP, and DSA topics.
5. Add measurable achievements to resume.""")

    add_heading(doc, "28. How to Run Locally")
    code_block(doc, """git clone https://github.com/07Jatin/Collage-placement-readiness-and-gap-analysis-testing-1-
cd "final project"

# Backend
python -m venv .venv
.venv\\Scripts\\activate
pip install -r requirements.txt
cd backend
uvicorn main:app --reload --port 8000

# Frontend
cd ../frontend
npm install
npm start

# Electron app from project root
npm install
npm run dev""")

    add_heading(doc, "29. Common Errors and Fixes")
    add_table(doc, ["Error", "Likely Fix"], [
        ("npm install fails", "Clear npm cache, delete broken node_modules if needed, then reinstall."),
        ("Port already in use", "Stop the existing process or change the frontend/backend port."),
        ("Backend not reachable", "Confirm FastAPI is running on http://127.0.0.1:8000 and CORS allows the frontend origin."),
        ("External API returns error", "Check API keys, quotas, internet access, and fallback behavior."),
        ("Resume parsing weak", "Improve keyword dictionaries and enable the LLM parser when configured."),
        ("Admin route forbidden", "Send Authorization: Bearer <token> after successful admin login."),
    ], [2.0, 4.4])

    add_heading(doc, "30. Summary")
    add_para(doc, "Placify AI helps students and colleges evaluate placement preparation in a structured way. It collects student profile data, assesses skills, analyzes resumes, calculates readiness, identifies missing skills, and recommends improvement actions.")
    code_block(doc, "Student Data + Assessments + Resume Evidence + Market Skills = Readiness Score + Gap Analysis + Recommendations")

    add_heading(doc, "31. Suggested README Structure")
    code_block(doc, """# Placify AI - Placement Readiness and Gap Analysis

## Introduction
Placify AI helps students analyze placement readiness, identify skill gaps, and follow personalized learning paths.

## Features
- Student and admin login
- Student profile management
- Resume skill extraction
- Skill quizzes and mock tests
- Placement readiness score
- Gap analysis report
- Personalized recommendations
- Admin dashboard and at-risk student tracking
- DSA coding lab

## Tech Stack
- Frontend: React, TailwindCSS, Recharts
- Backend: FastAPI, Python
- Data: JSON and CSV
- Desktop: Electron

## Installation
See the local setup section in the project documentation.""")

    add_heading(doc, "32. Final Beginner Understanding")
    add_numbered(doc, [
        "Understand the placement-readiness problem.",
        "Identify the users: student, admin/faculty, and placement cell.",
        "Learn what student data is collected.",
        "Understand how skills, resumes, tests, and target roles influence readiness.",
        "Study the backend APIs and service modules.",
        "Study the React views and how they call APIs.",
        "Understand how local JSON/CSV files store project data.",
        "Review the gap analysis and recommendation logic.",
        "Test the project end to end.",
        "Plan future improvements such as database migration and AI-based learning plans.",
    ])

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = footer.add_run("Placify AI - College Placement Readiness and Gap Analysis")
    rr.font.size = Pt(8)
    rr.font.color.rgb = RGBColor.from_string(GRAY)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    out = build_doc()
    print(out)
